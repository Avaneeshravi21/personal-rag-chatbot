"""
Loaders that turn raw files (pdf/docx/md/txt) into a common
Document representation: {text, metadata}.

Design note for your resume/interview: keeping this as one common
interface (`Document`) means the rest of the pipeline (chunker,
embedder, vector store) never needs to know what file format the
text originally came from.
"""
import os
import re
from dataclasses import dataclass, field
from typing import Optional
from pypdf import PdfReader
import docx


def _normalize_whitespace(text: str) -> str:
    """
    Fixes a common pypdf artifact on multi-column academic PDFs where
    extraction puts every word on its own line (e.g. "Retrieval\nAugmented\n
    Generation" instead of "Retrieval Augmented Generation"). We collapse
    single newlines into spaces but preserve real paragraph breaks (blank
    lines / double newlines), so recursive chunking still splits sensibly.
    """
    # Preserve paragraph breaks: mark blank-line-separated breaks first
    text = re.sub(r"\n\s*\n", "<<PARA>>", text)
    # Collapse remaining single newlines (the word-per-line artifact) into spaces
    text = text.replace("\n", " ")
    # Restore paragraph breaks
    text = text.replace("<<PARA>>", "\n\n")
    # Collapse repeated spaces
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


@dataclass
class Document:
    text: str
    source: str                 # filename
    doc_type: str                # "pdf" | "docx" | "md" | "txt"
    page: Optional[int] = None   # page number if applicable
    metadata: dict = field(default_factory=dict)


def load_pdf(filepath: str) -> list[Document]:
    """One Document per page — keeps page-level citations possible later."""
    reader = PdfReader(filepath)
    filename = os.path.basename(filepath)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = _normalize_whitespace(text)
        if text.strip():
            docs.append(
                Document(
                    text=text,
                    source=filename,
                    doc_type="pdf",
                    page=i + 1,
                )
            )
    return docs


def load_docx(filepath: str) -> list[Document]:
    d = docx.Document(filepath)
    filename = os.path.basename(filepath)
    full_text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
    full_text = _normalize_whitespace(full_text)
    return [Document(text=full_text, source=filename, doc_type="docx")]


def load_text_or_markdown(filepath: str) -> list[Document]:
    filename = os.path.basename(filepath)
    ext = "md" if filepath.endswith(".md") else "txt"
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(text=text, source=filename, doc_type=ext)]


LOADER_MAP = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".md": load_text_or_markdown,
    ".txt": load_text_or_markdown,
}


def load_directory(dir_path: str) -> list[Document]:
    """Walks a directory and loads every supported file into Documents."""
    all_docs: list[Document] = []
    if not os.path.isdir(dir_path):
        print(f"[warn] directory not found: {dir_path}")
        return all_docs

    for root, _, files in os.walk(dir_path):
        for fname in files:
            ext = os.path.splitext(fname)[1].lower()
            loader = LOADER_MAP.get(ext)
            if not loader:
                continue
            filepath = os.path.join(root, fname)
            try:
                docs = loader(filepath)
                all_docs.extend(docs)
                print(f"[loaded] {fname} -> {len(docs)} doc(s)")
            except Exception as e:
                print(f"[error] failed to load {fname}: {e}")

    return all_docs


if __name__ == "__main__":
    # quick manual test: point at data/raw/personal_docs
    import sys
    target = sys.argv[1] if len(sys.argv) > 1 else "data/raw/personal_docs"
    docs = load_directory(target)
    print(f"\nTotal documents loaded: {len(docs)}")
    for d in docs[:3]:
        print(f"- {d.source} (page {d.page}): {d.text[:100]!r}")
